from docx import Document
import os

import sys

if len(sys.argv) > 1:
    doc_path = sys.argv[1]
    # Auto-generate output path by replacing .docx with .txt
    output_path = doc_path.rsplit('.', 1)[0] + '.txt'
else:
    doc_path = r'c:\Users\nguye\OneDrive\Desktop\Chatbot_Crag\information\DATN_CHATBOT_RAG.docx'
    output_path = r'c:\Users\nguye\OneDrive\Desktop\Chatbot_Crag\information\DATN_CHATBOT_RAG.txt'

doc = Document(doc_path)
content = []

for para in doc.paragraphs:
    text = para.text.strip()
    if text:
        content.append(text)

# Also extract tables
for table in doc.tables:
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        content.append(" | ".join(cells))

with open(output_path, 'w', encoding='utf-8') as f:
    f.write('\n'.join(content))

print(f"Exported to {output_path}")
print(f"Total paragraphs: {len(content)}")
